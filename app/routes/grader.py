# AUTO-EXTRACTED grader routes from lti_routes.py
# Drop this file at: app/routes/grader.py
# Note: You may need to move helper functions/constants imported in lti_routes.py into utils and import them here.

# app/routes/grader.py  — imports
import json
import os
import shutil
import uuid
from datetime import datetime, timedelta

import requests
from docx import Document
from flask import (
    current_app,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from requests_oauthlib import OAuth1Session
from werkzeug.utils import secure_filename
from app.utils.slug import slugify
from app.supabase_client import supabase, upload_to_supabase
from app.utils.assignment_resolver import resolve_assignment_from_launch


from ..launch_utils import load_assignment_config
from ..utils.ai_usage_logger import log_ai_usage  # if you log usage
from ..utils.auth_decorators import require_tool
from ..utils.extractor import extract_filled_fields_from_pdf, extract_pdf_text
from ..utils.gpt_logging import log_gpt_interaction
from ..utils.grading_functions import (
    compare_fields_i130a,
    compare_fields_i765,
    compare_fields_n400,
)
from ..utils.text_utils import normalize_title

# 🔁 switch to relative imports so the package name doesn't matter
from . import lti

# --- Feature flag (optional) ---
FERPA_SAFE_MODE = (os.getenv("FERPA_SAFE_MODE") or "false").strip().lower() in {
    "1",
    "true",
    "yes",
    "on",
}


# --- RLS helper: single source + global hook ---
from uuid import UUID

ALL_ZERO_UUID = "00000000-0000-0000-0000-000000000000"


def apply_rls_uid():
    """Tell Supabase RLS who the caller is (idempotent)."""
    try:
        if not supabase:
            return

        # Prefer a real UUID when available; otherwise fall back to a DEV-safe UUID
        raw_uid = (
            ALL_ZERO_UUID
            if (session.get("is_superuser") or session.get("role") == "superuser")
            else (session.get("user_id") or session.get("student_id"))
        )

        if not raw_uid:
            return

        DEV_FAKE_UID = os.getenv("DEV_FAKE_UID", "00000000-0000-0000-0000-000000000001")

        try:
            # If this succeeds, raw_uid is a UUID and we can use it directly
            UUID(str(raw_uid))
            effective_uid = str(raw_uid)
        except Exception:
            # Not a UUID (e.g., "demo_user", "byu"); use a deterministic safe UUID so RLS passes
            effective_uid = DEV_FAKE_UID

        supabase.rpc("set_client_uid", {"uid": effective_uid}).execute()
    except Exception as e:
        try:
            current_app.logger.info(f"apply_rls_uid fallback/skip: {e}")
        except Exception:
            print("apply_rls_uid fallback/skip:", e)

# --- Back-compat alias so older links to /grader-base still work ---
@lti.route("/grader-base", methods=["GET"])
def _grader_base_alias():
    from flask import redirect, url_for
    return redirect(url_for("lti.grader_base"), code=302)

@lti.route("/_supabase-health")
def _supabase_health():
    import base64
    import json
    import os

    out = {}

    def _b64url_decode(s):
        s = s + "=" * (-len(s) % 4)
        return base64.urlsafe_b64decode(s.encode())

    try:
        # --- Show env pointing & JWT role (accurate) ---
        key = os.getenv("SUPABASE_KEY") or ""
        url = os.getenv("SUPABASE_URL") or ""
        proj = os.getenv("SUPABASE_PROJECT_ID") or ""

        out["SUPABASE_URL"] = url
        out["SUPABASE_PROJECT_ID"] = proj

        # Decode JWT payload without verifying signature
        jwt_role = "unknown"
        try:
            parts = key.split(".")
            if len(parts) >= 2:
                payload = json.loads(_b64url_decode(parts[1]).decode("utf-8"))
                jwt_role = (
                    payload.get("role")
                    or payload.get("app_metadata", {}).get("role")
                    or "unknown"
                )
                out["JWT_role_claim"] = jwt_role
            else:
                out["JWT_role_claim"] = "not-a-jwt"
        except Exception as e:
            out["JWT_role_claim"] = f"decode-failed: {e}"

        out["SUPABASE_KEY_is_service_role"] = jwt_role == "service_role"

        # --- RPC probe (RLS impersonation hook) ---
        try:
            supabase.rpc(
                "set_client_uid", {"uid": "00000000-0000-0000-0000-000000000001"}
            ).execute()
            out["rpc_set_client_uid_ok"] = True
        except Exception as e:
            out["rpc_set_client_uid_ok"] = False
            out["rpc_set_client_uid_err"] = str(e)

        # --- Minimal, RLS-sensitive reads (no writes) ---
        try:
            r = (
                supabase.table("assignments")
                .select("assignment_id,tool", count="exact")
                .eq("tool", "grader")
                .limit(1)
                .execute()
            )
            out["assignments_grader_count_exact"] = r.count
        except Exception as e:
            out["assignments_grader_count_err"] = str(e)

        try:
            r2 = (
                supabase.table("uscis_assignments")
                .select("assignment_id", count="exact")
                .limit(1)
                .execute()
            )
            out["uscis_assignments_count_exact"] = r2.count
        except Exception as e:
            out["uscis_assignments_count_err"] = str(e)

        try:
            r3 = (
                supabase.table("submissions")
                .select("submission_id,tool", count="exact")
                .eq("tool", "grader")
                .limit(1)
                .execute()
            )
            out["submissions_grader_count_exact"] = r3.count
        except Exception as e:
            out["submissions_grader_count_err"] = str(e)

    except Exception as e:
        out["fatal"] = str(e)
        out["traceback"] = traceback.format_exc()

    return "\n".join([f"{k}: {v}" for k, v in out.items()])


@lti.before_app_request
def _grader_rls_hook():
    # Runs before every request handled by this blueprint
    apply_rls_uid()


@lti.route("/grader", methods=["GET"], endpoint="grader_base")
@require_tool("grader")
def grader_base():
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    # 🔐 Tell PostgREST who we are for RLS policies
    apply_rls_uid()

    inst_id = session.get("institution_id")
    course_id = session.get("course_id")

    try:
        if session.get("is_superuser"):
            print("👑 Superuser — loading all Grader assignments")
            resp = (
                supabase.table("assignments")
                .select(
                    "assignment_id, assignment_title, tool, created_at, institution_id, course_id"
                )
                .eq("tool", "grader")  # only show Grader items
                .order("created_at", desc=True)
                .limit(300)
                .execute()
            )
        else:
            print("👤 Instructor — filter by institution/course (allow legacy NULLs)")
            q = (
                supabase.table("assignments")
                .select(
                    "assignment_id, assignment_title, tool, created_at, institution_id, course_id"
                )
                .eq("tool", "grader")
            )

            # Only add these OR-filters if the session value exists.
            if inst_id:
                q = q.or_(f"institution_id.eq.{inst_id},institution_id.is.null")
            if course_id:
                q = q.or_(f"course_id.eq.{course_id},course_id.is.null")

            resp = q.order("created_at", desc=True).limit(300).execute()

        assignments = resp.data or []
    except Exception as e:
        print("❌ Supabase fetch error in /grader-base:", repr(e))
        assignments = []

    return render_template("grader/grader_base.html", assignments=assignments)


@lti.route("/student-demo", methods=["GET"])
def student_demo_iframe():
    assignment_title = request.args.get("title", "").strip().lower()
    session["tool_role"] = "student"
    session["student_id"] = "demo_student"
    session["user_id"] = "demo_user"

    print("✅ [student-demo] session['user_id'] =", session.get("user_id"))

    session["platform"] = "demo"
    session["course_id"] = "demo_course"

    session["launch_data"] = {
        "https://purl.imsglobal.org/spec/lti/claim/resource_link": {
            "title": assignment_title
        },
        "https://purl.imsglobal.org/spec/lti/claim/roles": ["Student"],
        "given_name": "Demo User",
    }

    assignment_config = load_assignment_config(assignment_title)

    # 🚑 Fallback for iframe-based USCIS access
    if not assignment_config:
        if assignment_title == "form i-130a":
            assignment_config = {"form_type": "i130a"}
        elif assignment_title == "form n-400":
            assignment_config = {"form_type": "n400"}
        elif assignment_title == "form i-765":
            assignment_config = {"form_type": "i765"}
        else:
            assignment_config = {}

    return render_template(
        "launch.html",
        activity_name=assignment_title,
        user_name="Demo User",
        user_roles=["Student"],
        assignment_config=assignment_config,
        tinymce_api_key=os.getenv("TINYMCE_API_KEY"),
    )


@lti.route("/grade-docx", methods=["POST"])
def grade_docx():
    print("Superuser session flag:", session.get("is_superuser"))
    print("🎯 Reached grade-docx")
    print("🚨 /grade-docx route HIT")
    print("🔎 DEBUG ROUTE VERSION: Aug 20 — unified writer, safe UUIDs")
    print(f"🔐 FERPA_SAFE_MODE: {FERPA_SAFE_MODE}")

    # --- Local imports used in this route ---
    import json
    import os
    import re
    import uuid
    from datetime import datetime, timedelta
    from io import BytesIO

    import openai
    from werkzeug.utils import secure_filename

    # -# --- RESOLVE: slug > custom > title ---
    launch_data = session.get("launch_data", {}) or {}
    assignment_row, resolved_title, resolved_slug = resolve_assignment_from_launch(launch_data, request)

    if not assignment_row:
        return "❌ Assignment not found. Please contact your instructor.", 400

    assignment_title = resolved_title or (assignment_row.get("assignment_title") or "")
    assignment_config = assignment_row  # use DB row as the single source of truth
    print("🎯 Resolved assignment — title:", assignment_title, "| slug:", resolved_slug)


    # 🔎 Resolve assignment_id (some schemas require NOT NULL / FK)
    assignment_id_db = None
    try:
        aresp = (
            supabase.table("assignments")
            .select("assignment_id, tool")
            .eq("assignment_title", assignment_title)
            .single()
            .execute()
        )
        if aresp and aresp.data:
            assignment_id_db = aresp.data.get("assignment_id")
            if (aresp.data.get("tool") or "grader").lower() != "grader":
                print("⚠️ Assignment tool mismatch; continuing:", aresp.data.get("tool"))
    except Exception as e:
        print("ℹ️ Could not resolve assignment_id:", repr(e))

    gpt_model = assignment_config.get("gpt_model", "gpt-4")
    delay_setting = assignment_config.get("delay_posting", "immediate")
    delay_hours = {
        "immediate": 0,
        "1m": 0.0166,
        "12h": 12,
        "24h": 24,
        "36h": 36,
        "48h": 48,
    }.get(delay_setting, 0)

    rubric_url = assignment_config.get("rubric_file", "")
    file = request.files.get("file")
    inline_text = (request.form.get("inline_text") or "").strip()

    print("📎 Uploaded file object:", file)
    print("📝 Inline text received:", inline_text)

    if (not file or file.filename.strip() == "") and not inline_text:
        return "❌ No submission detected. Please upload a file or enter text.", 400

    # ---------- Init grading vars ----------
    file_ext = ""
    full_text = ""
    student_file_url = None
    score = 0
    total = 0
    feedback = ""
    incorrect_fields = []
    rubric_total_points = assignment_config.get("total_points", 100)

    # ---------- File upload + text extraction ----------
    if file:
        filename = file.filename.lower()
        file_ext = os.path.splitext(filename)[-1]
        safe_name = secure_filename(filename)
        unique_path = f"{uuid.uuid4()}_{safe_name}"
        file_bytes = file.read()

        # Upload original file to Storage (for review preview)
        try:
            supabase.storage.from_("submissions").upload(unique_path, file_bytes)
            SUPABASE_PROJECT_ID = os.getenv("SUPABASE_PROJECT_ID")
            student_file_url = f"https://{SUPABASE_PROJECT_ID}/storage/v1/object/public/submissions/{unique_path}"
            print("📎 Uploaded submission file to:", student_file_url)
        except Exception as e:
            print(
                "⚠️ Upload to submissions bucket failed (continuing without preview):",
                str(e),
            )
            student_file_url = None

        try:
            if gpt_model == "json":
                # === NoMas / Answer-key JSON mode (expects PDF) ===
                if file_ext != ".pdf":
                    return (
                        "❌ Grading mode is set to Answer Key (JSON), but file is not a PDF.",
                        400,
                    )

                raw_fields = extract_filled_fields_from_pdf(BytesIO(file_bytes))
                with open("data/debug_n400_extracted_fields.json", "w") as f:
                    json.dump(raw_fields, f, indent=2)
                print("🔬 DEBUG — Extracted sample keys:", list(raw_fields.keys())[:20])

                answer_key_url = assignment_config.get("answer_key_file") or rubric_url
                if not answer_key_url:
                    return "❌ No answer key found for this assignment.", 400

                resp = requests.get(answer_key_url)
                resp.raise_for_status()
                answer_key = resp.json()
                print("📘 Answer key loaded")

                # Compare helper supports either flat dict or {sections:[{fields:[]}]}
                def compare_fields(student_data, answer_key):
                    feedback_lines, incorrect_keys = [], []
                    sc, tot = 0, 0
                    fields = []

                    if isinstance(answer_key, dict) and "sections" in answer_key:
                        fields = [
                            f
                            for sec in answer_key["sections"]
                            for f in sec.get("fields", [])
                        ]
                    else:
                        for k, v in answer_key.items():
                            if isinstance(v, dict):
                                fields.append(
                                    {
                                        "field": k,
                                        "expected": v.get("expected", ""),
                                        "match": v.get("match", "exact"),
                                    }
                                )
                            else:
                                fields.append(
                                    {"field": k, "expected": v, "match": "exact"}
                                )

                    for fld in fields:
                        key = fld["field"]
                        expected = str(fld.get("expected", "")).strip().lower()
                        match = fld.get("match", "exact")
                        tot += 1

                        val = str(student_data.get(key, "") or "").strip().lower()
                        if val == "off":
                            base = key.rsplit("_", 1)[0]
                            candidates = {
                                k: v
                                for k, v in student_data.items()
                                if k.startswith(base) and str(v).lower() != "off"
                            }
                            val = (
                                str(list(candidates.values())[0]).strip().lower()
                                if candidates
                                else ""
                            )

                        if not val:
                            base = key.rsplit("_", 1)[0]
                            val = (
                                str(
                                    next(
                                        (
                                            v
                                            for k2, v in student_data.items()
                                            if k2.startswith(base)
                                            and str(v).lower() != "off"
                                        ),
                                        "",
                                    )
                                )
                                .strip()
                                .lower()
                            )

                        if not val:
                            incorrect_keys.append(key)
                            feedback_lines.append(
                                f"⚠️ Field '{key}' is empty or missing."
                            )
                        elif match == "exact" and val != expected:
                            incorrect_keys.append(key)
                            feedback_lines.append(
                                f"❌ Field '{key}' appears incorrect. Expected '{expected}' but got '{val}'."
                            )
                        else:
                            sc += 1

                    feedback_lines.append(
                        f"\nYou have {len(incorrect_keys)} error{'s' if len(incorrect_keys) != 1 else ''} in your submission."
                    )
                    feedback_lines.append(f"\n✅ Score: {sc} / {tot}")
                    return {
                        "score": sc,
                        "total": tot,
                        "feedback": "\n".join(feedback_lines),
                        "incorrect_fields": incorrect_keys,
                    }

                result = compare_fields(raw_fields, answer_key)
                score = result["score"]
                total = result["total"]
                feedback = result["feedback"]
                incorrect_fields = result.get("incorrect_fields", [])
                full_text = json.dumps(raw_fields, indent=2)

                with open(os.path.join("data", "last_mapped_fields.json"), "w") as f:
                    f.write(full_text)
                print("✅ Saved mapped fields to data/last_mapped_fields.json")

            else:
                # === Regular Rubiqs Grader mode ===
                if file_ext == ".pdf":
                    full_text = extract_pdf_text(BytesIO(file_bytes))
                elif file_ext == ".docx":
                    from docx import Document

                    doc = Document(BytesIO(file_bytes))
                    full_text = "\n".join(
                        [p.text for p in doc.paragraphs if p.text.strip()]
                    )
                else:
                    return "❌ Unsupported file type. Please upload .docx or .pdf", 400

        except Exception:
            print("❌ Critical grading failure:")
            import traceback

            traceback.print_exc()
            return render_template(
                "feedback.html",
                pending_message="❌ Something went wrong while processing your submission. Please try again or contact your instructor.",
            )

    elif inline_text:
        full_text = inline_text

    # ---------- GPT rubric scoring (non-JSON mode) ----------
    if gpt_model != "json":
        try:
            r = requests.get(rubric_url)
            print("🌐 Downloaded rubric file:", rubric_url, "status:", r.status_code)
            if r.status_code != 200:
                return f"❌ Failed to download rubric file. Status {r.status_code}", 500

            rubric_content = r.content
            rubric_json = None
            rubric_text = ""

            if rubric_url.endswith(".json"):
                rubric_json = json.loads(rubric_content)
                if "criteria" in rubric_json:
                    rubric_text = "\n".join(
                        [f"- {c['description']}" for c in rubric_json["criteria"]]
                    )
                    rubric_total_points = sum(
                        c.get("max_points", 1) for c in rubric_json["criteria"]
                    )
                elif "sections" in rubric_json:
                    rubric_text = "\n".join(
                        [
                            f"- {f['field']}: expected '{f['expected']}'"
                            for s in rubric_json["sections"]
                            for f in s.get("fields", [])
                        ]
                    )
                    rubric_total_points = assignment_config.get("total_points", 10)
                else:
                    rubric_text = "(Invalid JSON rubric format)"
                    rubric_total_points = assignment_config.get("total_points", 100)
            elif rubric_url.endswith(".docx"):
                from docx import Document

                doc = Document(BytesIO(rubric_content))
                rubric_text = "\n".join([p.text for p in doc.paragraphs])
                rubric_total_points = assignment_config.get("total_points", 100)
            elif rubric_url.endswith(".pdf"):
                rubric_text = extract_pdf_text(BytesIO(rubric_content))
                rubric_total_points = assignment_config.get("total_points", 100)
            else:
                rubric_text = "(Unknown rubric format)"
                rubric_total_points = assignment_config.get("total_points", 100)

            grading_difficulty = assignment_config.get("grading_difficulty", "balanced")
            student_level = assignment_config.get("student_level", "college")
            feedback_tone = assignment_config.get("feedback_tone", "supportive")
            ai_notes = assignment_config.get("ai_notes", "")

            prompt = f"""
You are a helpful AI grader.

Assignment Title: {assignment_title}
Grading Difficulty: {grading_difficulty}
Student Level: {student_level}
Feedback Tone: {feedback_tone}
Total Points: {rubric_total_points}

Rubric:
{rubric_text[:2000]}
"""
            if ai_notes:
                prompt += f"\nInstructor Notes:\n{ai_notes}"

            prompt += f"""

Student Submission:
---
{full_text[:3000]}
---

Return your response in this format:

Score: <number from 0 to {rubric_total_points}>
Feedback: <detailed, helpful feedback>
""".strip()

            openai.api_key = os.getenv("OPENAI_API_KEY")
            resp = openai.ChatCompletion.create(
                model=gpt_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=1000,
            )
            output = resp["choices"][0]["message"]["content"]
            usage = resp.get("usage", {})

            # Billing/reporting
            log_ai_usage(
                user_id=session.get("user_id"),
                institution_id=session.get("institution_id"),
                tool="grader",
                model=gpt_model,
                assignment_id=assignment_title,
                usage=usage,
            )

            m = re.search(r"Score:\s*(\d{1,3})", output)
            score = int(m.group(1)) if m else 0
            fm = re.search(r"Feedback:\s*(.+)", output, re.DOTALL)
            feedback = fm.group(1).strip() if fm else output.strip()

        except openai.error.OpenAIError as e:
            return f"❌ GPT error: {str(e)}", 500
        except Exception as e:
            return f"❌ Rubric or prompt error: {str(e)}", 500

    # ---------- RLS helper (safe with non-UUID dev ids) ----------
    if not session.get("student_id"):
        session["student_id"] = session.get("launch_data", {}).get("sub")

    def _is_uuid(v):
        try:
            uuid.UUID(str(v))
            return True
        except Exception:
            return False

    _uid = session.get("student_id") or session.get("user_id")
    try:
        if _is_uuid(_uid):
            supabase.rpc("set_client_uid", {"uid": str(_uid)}).execute()
            print("🔐 set_client_uid ->", _uid)
        else:
            # Use a deterministic fake UUID for dev so RLS sees a non-null UID
            DEV_FAKE_UID = os.getenv(
                "DEV_FAKE_UID", "00000000-0000-0000-0000-000000000001"
            )
            supabase.rpc("set_client_uid", {"uid": DEV_FAKE_UID}).execute()
            print("🔐 set_client_uid -> DEV_FAKE_UID", DEV_FAKE_UID)

    except Exception as e:
        print("⚠️ set_client_uid RPC failed (continuing):", str(e))

    # ---------- Build submission payload ----------
    submission_id = str(uuid.uuid4())
    now = datetime.utcnow()
    release_time = now + timedelta(hours=delay_hours)
    ready_to_post = delay_hours == 0 and not assignment_config.get(
        "instructor_approval", False
    )

    submission_data = {
        "submission_id": submission_id,
        "student_id": session["student_id"],
        "assignment_title": assignment_title,
        "assignment_id": assignment_id_db,
        "course_id": session.get("course_id", "demo_course"),
        "institution_id": session.get("institution_id"),
        "submission_time": now.isoformat() + "Z",
        "tool": "grader",
        "score": score,
        "feedback": feedback,
        "submission_type": "inline" if inline_text else "file",
        "student_text": full_text,
        "student_file_url": student_file_url,
        "ai_check_result": None,
        "instructor_notes": "",
        "delay_hours": delay_hours,
        "ready_to_post": ready_to_post,
        "pending": not ready_to_post,
        "reviewed": False,
        "release_time": release_time.isoformat(),
        "incorrect_fields": incorrect_fields if gpt_model == "json" else [],
    }

    # --- Optional: force Grader (non-JSON mode) into pending review queue by default ---
    if (assignment_config.get("gpt_model") or "").lower() != "json":
        submission_data["ready_to_post"] = False
        submission_data["pending"] = True

    # ---------- Single, unified writer (NO duplicates) ----------
    try:
        print("✅ Submitting to Supabase:", submission_data["submission_id"])

        # Detect NoMas/USCIS by DB title or config hint
        is_nomas = False
        resolved_form_type = None
        arow = {}

        # A) Exact title match in uscis_assignments
        try:
            aresp = (
                supabase.table("uscis_assignments")
                .select("assignment_id, assignment_title, form_type")
                .eq("assignment_title", assignment_title)
                .single()
                .execute()
            )
            arow = aresp.data or {}
        except Exception:
            arow = {}

        # B) Fallback ilike
        if not arow:
            try:
                aresp2 = (
                    supabase.table("uscis_assignments")
                    .select("assignment_id, assignment_title, form_type")
                    .ilike("assignment_title", f"%{assignment_title}%")
                    .limit(1)
                    .execute()
                )
                if aresp2.data:
                    arow = aresp2.data[0]
            except Exception:
                pass

        if arow:
            is_nomas = True
            resolved_form_type = (arow.get("form_type") or "").lower()
        else:
            # C) Config hints (form-type assignment or JSON mode implies NoMas)
            cfg_ft = (assignment_config or {}).get("form_type")
            cfg_type = (assignment_config or {}).get("assignment_type")
            if (
                (cfg_ft and cfg_ft.strip())
                or (cfg_type and cfg_type.lower() == "uscis")
                or (gpt_model == "json")
            ):
                is_nomas = True
                resolved_form_type = (cfg_ft or "").lower()

        if is_nomas:
            # uscis_submissions: student_id is TEXT → safe for demo strings
            payload = {
                "submission_id": submission_data["submission_id"],
                "student_id": submission_data["student_id"],
                "assignment_title": submission_data["assignment_title"],
                "form_type": resolved_form_type,
                "submission_time": submission_data["submission_time"],
                "submitted_at": submission_data["submission_time"],
                "score": submission_data["score"],
                "total": total,
                "feedback": submission_data["feedback"],
                "incorrect_fields": submission_data.get("incorrect_fields", []),
                "student_file_url": submission_data["student_file_url"],
                "student_text": submission_data["student_text"],
                "delay_hours": submission_data["delay_hours"],
                "release_time": submission_data["release_time"],
                "ready_to_post": submission_data["ready_to_post"],
                "pending": submission_data["pending"],
                "reviewed": submission_data["reviewed"],
                "instructor_notes": submission_data["instructor_notes"],
            }
            supabase.table("uscis_submissions").insert(payload).execute()
            print("🗄️ Wrote to uscis_submissions")
        else:
            # === Generic Rubiqs Grader submission -> public.submissions (full, RLS-safe) ===

            # Use the same effective UID you set for RLS so row passes policies
            effective_uid = (
                _uid
                if _is_uuid(_uid)
                else os.getenv("DEV_FAKE_UID", "00000000-0000-0000-0000-000000000001")
            )

            # NEW: legacy text id for NOT NULL column student_id_text_old
            legacy_sid_text = str(
                session.get("student_id") or session.get("user_id") or effective_uid
            )

            # Ensure the payload we insert includes the columns your UI & RLS need
            row = {
                # identifiers / scoping
                "submission_id": submission_data["submission_id"],
                "tool": "grader",
                "student_id": effective_uid,
                "institution_id": session.get("institution_id"),
                "course_id": session.get("course_id", "demo_course"),
                # NEW: satisfy NOT NULL legacy column
                "student_id_text_old": legacy_sid_text,
                # assignment + timing
                "assignment_title": assignment_title,
                "submission_time": now.isoformat() + "Z",
                "release_time": submission_data["release_time"],
                # grading status
                "score": score,
                "feedback": feedback,
                "pending": submission_data["pending"],
                "reviewed": submission_data["reviewed"],
                "ready_to_post": submission_data["ready_to_post"],
                # content (only if these columns exist)
                "submission_type": submission_data["submission_type"],
                "student_text": submission_data["student_text"],
                "student_file_url": submission_data["student_file_url"],
            }

            try:
                resp = supabase.table("submissions").insert(row).execute()
                if getattr(resp, "error", None):
                    print("❌ Supabase insert error:", resp.error)

                    # Fallback minimal row also needs legacy column:
                    minimal_row = {
                        "submission_id": row["submission_id"],
                        "tool": "grader",
                        "student_id": row["student_id"],
                        "student_id_text_old": legacy_sid_text,  # <-- NEW here too
                        "assignment_title": row["assignment_title"],
                        "submission_time": row["submission_time"],
                        "pending": row["pending"],
                        "reviewed": row["reviewed"],
                        "ready_to_post": row["ready_to_post"],
                        "score": row["score"],
                        "feedback": row["feedback"],
                    }
                    print("↪️ Retrying with minimal_row:", minimal_row)
                    resp2 = supabase.table("submissions").insert(minimal_row).execute()
                    if getattr(resp2, "error", None):
                        print("💥 Supabase insert error (minimal row):", resp2.error)
                        return (
                            "❌ Failed to save your submission (DB error). Please contact your instructor.",
                            500,
                        )
                    saved2 = (resp2.data or [minimal_row])[0]
                    print(
                        "🗄️ Wrote to submissions (submission_id, minimal):",
                        saved2.get("submission_id") or saved2.get("id"),
                    )
                else:
                    saved = (resp.data or [row])[0]
                    print(
                        "🗄️ Wrote to submissions (submission_id):",
                        saved.get("submission_id") or saved.get("id"),
                    )
            except Exception as e2:
                print("💥 Insert to Supabase failed completely:", repr(e2))
                return (
                    "❌ Failed to save your submission (DB error). Please contact your instructor.",
                    500,
                )

        # Clean temp dirs if any were used
        shutil.rmtree("temp_uploads", ignore_errors=True)
        shutil.rmtree("converted", ignore_errors=True)
        shutil.rmtree("converted_images", ignore_errors=True)

    except Exception as e:
        import traceback

        print("💥 Insert to Supabase failed completely:", repr(e))
        traceback.print_exc()  # ⬅️ prints PostgREST error body / stack
        return (
            "❌ Failed to save your submission (DB error). Please contact your instructor.",
            500,
        )

    # ---------- Optional log ----------
    try:
        log_gpt_interaction(assignment_title, full_text, feedback, score)
    except Exception as e:
        print("❌ GPT log failed:", str(e))

    # ---------- Final response / AGS ----------
    if assignment_config.get("instructor_approval"):
        return render_template(
            "feedback.html",
            pending_message="✅ This submission requires instructor review. Your feedback will be posted after approval.",
        )
    elif delay_hours > 0:
        return render_template(
            "feedback.html",
            pending_message=f"⏳ This submission will be released after {delay_hours} hour(s).",
        )
    else:
        if session.get("platform") == "canvas":
            try:
                print("🎯 Canvas detected — attempting AGS post...")
                post_grade_to_lms(session, score, feedback)
            except Exception as e:
                print("❌ AGS post failed:", str(e))
        else:
            print(
                "ℹ️ AGS posting skipped — not Canvas (platform:",
                session.get("platform"),
                ")",
            )

        return render_template(
            "feedback.html",
            feedback=feedback,
            score=score if gpt_model != "json" else total - len(incorrect_fields),
            rubric_total_points=rubric_total_points if gpt_model != "json" else total,
            user_roles=session.get("launch_data", {}).get(
                "https://purl.imsglobal.org/spec/lti/claim/roles", []
            ),
        )


@lti.route("/nomas-dashboard", methods=["GET"], endpoint="nomas_dashboard")
def nomas_dashboard():
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    inst_id = session.get("institution_id")
    course_id = session.get("course_id", "demo_course")

    # --- Assignments ---
    try:
        if session.get("is_superuser"):
            aresp = (
                supabase.table("uscis_assignments")
                .select(
                    "assignment_id, assignment_title, form_type, institution_id, course_id, created_at"
                )
                .order("created_at", desc=True)
                .execute()
            )
            assignments = aresp.data or []
        else:
            aresp = (
                supabase.table("uscis_assignments")
                .select(
                    "assignment_id, assignment_title, form_type, institution_id, course_id, created_at"
                )
                .or_(f"institution_id.eq.{inst_id},institution_id.is.null")
                .or_(f"course_id.eq.{course_id},course_id.is.null")
                .order("created_at", desc=True)
                .execute()
            )
            assignments = aresp.data or []
    except Exception as e:
        print("❌ load uscis_assignments:", e)
        assignments = []

    # --- Submissions (order by submitted_at, fallback to submission_time) ---
    try:
        try:
            sresp = (
                supabase.table("uscis_submissions")
                .select("*")
                .order("submitted_at", desc=True)
                .limit(300)
                .execute()
            )
            submissions = sresp.data or []
        except Exception:
            sresp = (
                supabase.table("uscis_submissions")
                .select("*")
                .order("submission_time", desc=True)
                .limit(300)
                .execute()
            )
            submissions = sresp.data or []
    except Exception as e:
        print("❌ load uscis_submissions:", e)
        submissions = []

    # normalize aliases the template expects
    for s in submissions:
        s.setdefault("user_id", s.get("student_id"))
        s.setdefault("submitted_at", s.get("submission_time"))

    return render_template(
        "nomas_dashboard.html", assignments=assignments, submissions=submissions
    )


@lti.route("/uscis-dashboard")
def uscis_dashboard():
    course_id = session.get("course_id", "demo_course")

    # Filter USCIS assignments by course
    if session.get("is_superuser"):
        print("👑 Superuser: loading all USCIS assignments")
        assignments = (
            supabase.table("uscis_assignments").select("*").execute().data or []
        )
    else:
        print("👤 Instructor: filtering USCIS assignments by institution and course")
        assignments = (
            supabase.table("uscis_assignments")
            .select("*")
            .eq("institution_id", session.get("institution_id"))
            .eq("course_id", course_id)
            .execute()
            .data
            or []
        )

    return render_template("uscis_dashboard.html", assignments=assignments)


@lti.route("/generate-answer-key", methods=["POST"])
def generate_answer_key():
    # === Step 1: Read and validate file ===
    uploaded_file = request.files.get("file")
    if not uploaded_file or uploaded_file.filename == "":
        return "❌ No file uploaded.", 400

    if not uploaded_file.filename.lower().endswith(".pdf"):
        return "❌ Please upload a PDF file.", 400

    pdf_bytes = uploaded_file.read()
    raw_fields = extract_filled_fields_from_pdf(pdf_bytes)

    # === STEP 1.5: Normalize filled fields, keep only checked/filled ===
    filtered_fields = {}
    for key, value in raw_fields.items():
        value_normalized = str(value).strip().lower()
        if value_normalized in ["", "off"]:
            continue
        filtered_fields[key] = value

    # === STEP 1.6: Optional - Group known radio buttons under semantic labels ===
    grouped_fields = {}
    radio_groups = {
        "Reason for Filing": [
            "Page1_General Provision_1",
            "Page1_Spouse of U.S. Citizen_2",
            "Page1_VAWA_3",
        ]
    }

    for group_label, field_keys in radio_groups.items():
        for k in field_keys:
            if filtered_fields.get(k) == "On":
                grouped_fields[group_label] = {k: "Yes"}
                break

    # Add all non-radio fields back in
    for k, v in filtered_fields.items():
        if not any(k in group for group in radio_groups.values()):
            grouped_fields[k] = v

    fields = grouped_fields

    # === Step 2: Prepare Supabase client ===
    SUPABASE_URL = os.getenv("SUPABASE_URL")
    SUPABASE_KEY = os.getenv("SUPABASE_KEY")

    SUPABASE_PROJECT_ID = os.getenv("SUPABASE_PROJECT_ID")

    # === Step 3: Upload answer key as JSON to Supabase ===
    key_json = json.dumps(fields, indent=2)
    unique_filename = f"generated_keys/answer_key_{uuid.uuid4()}.json"

    try:
        supabase.storage.from_("rubrics").upload(
            unique_filename, key_json.encode("utf-8")
        )
    except Exception as e:
        return f"❌ Failed to upload to Supabase: {str(e)}", 500

    # === Step 4: Build public URL and redirect back to dashboard ===
    public_url = f"https://{SUPABASE_PROJECT_ID}/storage/v1/object/public/rubrics/{unique_filename}"
    return redirect(f"/nomas-dashboard?rubric_url={public_url}")


import re

UUID_RE = re.compile(
    r"^[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}$"
)


def _uuid_or_none(val):
    s = (val or "").strip()
    return s if UUID_RE.match(s) else None


@lti.route("/create-uscis-assignment", methods=["POST"])
def create_uscis_assignment():
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    print("🚨 create-uscis-assignment route HIT")

    SUPABASE_URL = os.getenv("SUPABASE_URL")
    SUPABASE_KEY = os.getenv("SUPABASE_KEY")

    SUPABASE_PROJECT_ID = os.getenv("SUPABASE_PROJECT_ID")

    # === Step 1: Pull all form data ===
    title = request.form.get("title", "").strip()
    form_type = request.form.get("form_type", "").strip().lower()

    rubric_url = request.form.get("rubric_file_url", "").strip()
    answer_key_file = request.files.get("answer_key_upload")

    if answer_key_file and answer_key_file.filename:
        try:
            filename = secure_filename(answer_key_file.filename)
            unique_path = f"generated_keys/{uuid.uuid4()}_{filename}"
            supabase.storage.from_("rubrics").upload(
                unique_path, answer_key_file.read()
            )
            SUPABASE_PROJECT_ID = os.getenv("SUPABASE_PROJECT_ID")
            rubric_url = f"https://{SUPABASE_PROJECT_ID}/storage/v1/object/public/rubrics/{unique_path}"
            print(f"✅ Uploaded answer key file to Supabase: {rubric_url}")
        except Exception as e:
            print("❌ Error uploading answer key file:", str(e))

    require_review = request.form.get("require_review", "").lower() == "true"
    total_points = request.form.get("total_points")
    complete_incomplete = request.form.get("complete_incomplete") == "true"
    grade_level = request.form.get("grade_level")
    grading_difficulty = request.form.get("grading_difficulty")

    # Safely coerce potential UUIDs (omit if not UUID)
    created_by_uuid = _uuid_or_none(session.get("user_id"))
    institution_uuid = _uuid_or_none(session.get("institution_id"))

    assignment_data = {
        "assignment_id": f"a_{uuid.uuid4().hex[:8]}",
        "assignment_title": title,
        "form_type": form_type,
        "course_id": session.get("course_id", "demo_course"),
        "institution_id": session.get("institution_id"),
        "rubric_file": rubric_url,
        "answer_key_file": rubric_url,
        "gpt_model": "json",
        "instructor_approval": require_review,
        "total_points": int(total_points) if total_points else None,
        "complete_incomplete": complete_incomplete,
        "student_level": grade_level,
        "grading_difficulty": grading_difficulty,
        "created_at": datetime.utcnow().isoformat(),
        # "created_by": <removed — was causing UUID error when demo_user>
    }

    # Conditionally add UUID-typed columns ONLY if valid:
    if institution_uuid:
        assignment_data["institution_id"] = institution_uuid
    if created_by_uuid:
        assignment_data["created_by"] = created_by_uuid

    # === Step 3: Insert into Supabase ===
    try:
        supabase.table("uscis_assignments").insert(assignment_data).execute()
    except Exception as e:
        return f"❌ Failed to save assignment: {str(e)}", 500

    return redirect("/nomas-dashboard")


@lti.route("/grade-uscis-form", methods=["POST"])
def grade_uscis_form():
    print("🎯 Reached grade-uscis-form")
    print("🚨 /grade-uscis-form route HIT")
    print("🔎 DEBUG ROUTE VERSION: July 11 — USCIS forms only (N-400 for now)")
    print(f"🔐 FERPA_SAFE_MODE: {FERPA_SAFE_MODE}")

    # Local imports to keep this route self-contained
    import json
    import os
    import uuid
    from datetime import datetime
    from io import BytesIO
    from pprint import pprint

    import requests
    from werkzeug.utils import secure_filename

    # --- Build assignment_title from LTI claim ---
    resource_link = session.get("launch_data", {}).get(
        "https://purl.imsglobal.org/spec/lti/claim/resource_link", {}
    )
    title_from_claim = resource_link.get("title")
    id_from_claim = resource_link.get("id")
    assignment_title = normalize_title(
        str(
            title_from_claim or f"Assignment-{id_from_claim}" or "Untitled Assignment"
        ).strip()
    )
    print("🧠 resource_link:", resource_link)
    print("🧠 assignment_title after normalize:", assignment_title)

    # ✅ Data-driven: resolve assignment from Supabase first (no name guessing)
    assignment_config = {}
    try:
        aresp = (
            supabase.table("uscis_assignments")
            .select("*")
            .eq("assignment_title", assignment_title)
            .single()
            .execute()
        )
        assignment_config = aresp.data or {}
    except Exception as e:
        print("⚠️ uscis_assignments lookup failed:", str(e))

    # Fallback to legacy loader only if DB lookup failed
    if not assignment_config:
        assignment_config = load_assignment_config(assignment_title) or {}

    print("📦 assignment_config resolved:", assignment_config)

    form_type = (assignment_config.get("form_type") or "").lower()
    if not form_type:
        return (
            "❌ No form_type found for this assignment. "
            "Please ensure the assignment was created with a form type in NoMas.",
            400,
        )

    print("🎯 Resolved assignment title:", assignment_title)
    print("🧩 Using assignment_config:", assignment_config)

    # --- Validate uploaded file (PDF only for USCIS form grading) ---
    file = request.files.get("file")
    if not file or not file.filename.lower().endswith(".pdf"):
        return "❌ Please upload a valid PDF file.", 400

    # === Upload and Extract PDF Fields ===
    filename = secure_filename(file.filename.lower())
    file_bytes = file.read()
    unique_path = f"submissions/{uuid.uuid4()}_{filename}"

    try:
        raw_fields = extract_filled_fields_from_pdf(BytesIO(file_bytes))
        pprint(list(raw_fields.items())[:20])  # Debug: show sample fields

        # Save raw extraction for debugging
        with open("data/debug_extracted_fields.json", "w") as f:
            json.dump(raw_fields, f, indent=2)

        # Upload original PDF to Storage
        supabase.storage.from_("submissions").upload(unique_path, file_bytes)
        SUPABASE_PROJECT_ID = os.getenv("SUPABASE_PROJECT_ID")
        student_file_url = f"https://{SUPABASE_PROJECT_ID}/storage/v1/object/public/submissions/{unique_path}"
    except Exception as e:
        print("❌ PDF extraction failed:", str(e))
        return (
            "❌ Failed to process the uploaded form. Please contact your instructor.",
            500,
        )

    # === Load the answer key JSON from Supabase ===
    rubric_url = (
        assignment_config.get("answer_key_file")
        or assignment_config.get("rubric_file")
        or assignment_config.get("rubric_file_url")
    )
    print("📎 rubric_url resolved to:", rubric_url)

    if not rubric_url:
        return (
            "❌ This assignment is missing an answer key. Please contact your instructor.",
            400,
        )

    try:
        answer_key_response = requests.get(rubric_url)
        answer_key_response.raise_for_status()
        answer_key_json = answer_key_response.json()
    except Exception as e:
        print("❌ Failed to load answer key from URL:", rubric_url)
        print("📛 Error:", str(e))
        return "❌ Could not retrieve answer key. Please contact your instructor.", 500

    # === Compare fields according to form type ===
    if form_type == "n400":
        result = compare_fields_n400(raw_fields, answer_key_json)
    elif form_type == "i765":
        result = compare_fields_i765(raw_fields, answer_key_json)
    elif form_type == "i130a":
        result = compare_fields_i130a(raw_fields, answer_key_json)
    else:
        return jsonify({"error": f"❌ Unsupported form type '{form_type}'."}), 400

    score = result["score"]
    total = result["total"]
    feedback = result["feedback"]
    incorrect_fields = result["incorrect_fields"]

    full_text = json.dumps(raw_fields, indent=2)

    # Save last mapped/extracted fields for debugging
    debug_path = os.path.join("data", "last_mapped_fields.json")
    with open(debug_path, "w") as f:
        f.write(full_text)
    print(f"✅ Saved extracted fields to {debug_path}")

    # === Store Submission in Supabase ===
    if not session.get("student_id"):
        # Fall back to LTI subject if student_id not set yet
        session["student_id"] = session.get("launch_data", {}).get("sub")

    # Set client UID for RLS policies that rely on auth.uid() emulation
    try:
        supabase.rpc("set_client_uid", {"uid": session["student_id"]}).execute()
    except Exception as e:
        print("⚠️ set_client_uid RPC failed (continuing):", str(e))

    now = datetime.utcnow()
    delay_hours = {
        "immediate": 0,
        "1m": 0.0166,
        "12h": 12,
        "24h": 24,
        "36h": 36,
        "48h": 48,
    }.get(assignment_config.get("delay_posting", "immediate"), 0)
    release_time = now + timedelta(hours=delay_hours)
    ready_to_post = delay_hours == 0 and not assignment_config.get(
        "instructor_approval", False
    )

    submission_data = {
        "submission_id": str(uuid.uuid4()),
        "student_id": session["student_id"],
        "assignment_title": assignment_title,
        "course_id": session.get("course_id", "demo_course"),
        "institution_id": session.get("institution_id"),
        "submission_time": now.isoformat(),
        "score": score,
        "feedback": feedback,
        "submission_type": "file",
        "student_text": full_text,
        "student_file_url": student_file_url,
        "ai_check_result": None,
        "instructor_notes": "",
        "delay_hours": delay_hours,
        "ready_to_post": ready_to_post,
        "pending": not ready_to_post,
        "reviewed": False,
        "release_time": release_time.isoformat(),
        "incorrect_fields": incorrect_fields,
    }

    try:
        print("✅ Submitting to Supabase:", submission_data["submission_id"])

        # Insert ONLY columns that exist in uscis_submissions
        uscis_payload = {
            "submission_id": submission_data["submission_id"],
            "student_id": submission_data["student_id"],
            "assignment_title": submission_data["assignment_title"],
            "form_type": form_type,
            "submission_time": submission_data["submission_time"],
            "score": submission_data["score"],
            "total": total,  # <- include total from the compare result
            "feedback": submission_data["feedback"],
            "student_file_url": submission_data["student_file_url"],
            "student_text": submission_data["student_text"],
            "incorrect_fields": submission_data.get("incorrect_fields", []),
            "delay_hours": submission_data["delay_hours"],
            "release_time": submission_data["release_time"],
            "ready_to_post": submission_data["ready_to_post"],
            "pending": submission_data["pending"],
            "reviewed": submission_data["reviewed"],
            "instructor_notes": submission_data["instructor_notes"],
            # Optional: if you later add attempt_number, set it here (defaults ok if column has default)
            # "attempt_number": 1,
        }

        supabase.table("uscis_submissions").insert(uscis_payload).execute()
        print("🗄️ Wrote to uscis_submissions")

    except Exception as e:
        # Log the exact server error so you can see missing/invalid columns
        print("❌ uscis_submissions insert failed:", repr(e))
        try:
            # Safe fallback so you never lose a submission during the demo
            supabase.table("submissions").insert(submission_data).execute()
            print("↪️ Fallback: wrote to legacy submissions")
        except Exception as e2:
            print("❌ Legacy submissions insert also failed:", repr(e2))

    # Optional: log to your GPT interaction log (safe to continue if it fails)
    try:
        log_gpt_interaction(assignment_title, full_text, feedback, score)
    except Exception as e:
        print("❌ GPT log failed:", str(e))

    # === Display Feedback / AGS passback ===
    if assignment_config.get("instructor_approval"):
        return render_template(
            "feedback.html",
            pending_message="✅ Submission received. Awaiting instructor approval.",
        )
    elif delay_hours > 0:
        return render_template(
            "feedback.html",
            pending_message=f"⏳ Feedback will be released in {delay_hours} hour(s).",
        )
    else:
        if session.get("platform") == "canvas":
            try:
                post_grade_to_lms(session, score, feedback)
            except Exception as e:
                print("❌ Canvas grade passback failed:", str(e))
        return render_template(
            "feedback.html",
            feedback=feedback,
            score=score,
            rubric_total_points=total,  # total points for this form comparison
        )


@lti.route("/edit-uscis-assignment/<assignment_id>", methods=["GET", "POST"])
def edit_uscis_assignment(assignment_id):
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    if request.method == "POST":
        try:
            print("📥 USCIS Edit POST:", dict(request.form))

            # === Step 1: Load existing assignment ===
            existing = (
                supabase.table("uscis_assignments")
                .select("*")
                .eq("assignment_id", assignment_id)
                .limit(1)
                .execute()
                .data
            )
            if not existing:
                return f"❌ Assignment not found: {assignment_id}", 404
            assignment = existing[0]
            rubric_url = assignment.get("answer_key_file", "")

            # === Step 2: Check for uploaded answer key ===
            answer_key_file = request.files.get("answer_key_upload")
            if answer_key_file and answer_key_file.filename:
                import uuid

                from werkzeug.utils import secure_filename

                filename = secure_filename(answer_key_file.filename)
                unique_path = f"generated_keys/{uuid.uuid4()}_{filename}"
                supabase.storage.from_("rubrics").upload(
                    unique_path, answer_key_file.read()
                )
                SUPABASE_PROJECT_ID = os.getenv("SUPABASE_PROJECT_ID")
                rubric_url = f"https://{SUPABASE_PROJECT_ID}/storage/v1/object/public/rubrics/{unique_path}"
                print("✅ New answer key uploaded:", rubric_url)

            # === Step 3: Update fields ===
            updated_fields = {
                "assignment_title": request.form.get("title"),
                "form_type": request.form.get("form_type"),
                "answer_key_file": rubric_url,
                "rubric_file": rubric_url,
                "instructor_approval": request.form.get("require_review") == "true",
                "total_points": (
                    int(request.form.get("total_points"))
                    if request.form.get("total_points")
                    else None
                ),
                "complete_incomplete": request.form.get("complete_incomplete")
                == "true",
                "student_level": request.form.get("grade_level"),
                "grading_difficulty": request.form.get("grading_difficulty"),
            }

            supabase.table("uscis_assignments").update(updated_fields).eq(
                "assignment_id", assignment_id
            ).execute()
            print("✅ Assignment updated.")
            return redirect("/nomas-dashboard")

        except Exception as e:
            print("❌ Error updating USCIS assignment:", e)
            return f"Error: {e}", 500

    # === GET: Load current values ===
    response = (
        supabase.table("uscis_assignments")
        .select("*")
        .eq("assignment_id", assignment_id)
        .execute()
    )

    if not response.data:
        return f"❌ Assignment not found: {assignment_id}", 404

    assignment = response.data[0]
    return render_template("edit_uscis_assignment.html", assignment=assignment)


@lti.route("/delete-uscis-assignment", methods=["POST"])
def delete_uscis_assignment():
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    data = request.get_json()
    assignment_id = data.get("assignment_id")

    if not assignment_id:
        return jsonify({"success": False, "error": "Missing assignment_id"}), 400

    try:
        supabase.table("uscis_assignments").delete().eq(
            "assignment_id", assignment_id
        ).execute()
        return jsonify({"success": True})
    except Exception as e:
        print("❌ Error deleting USCIS assignment:", e)
        return jsonify({"success": False, "error": str(e)}), 500


@lti.route("/nomas/review", methods=["GET"], endpoint="nomas_review")
def nomas_review():
    sid = request.args.get("submission_id")
    if not sid:
        return redirect(url_for("lti.nomas_training_dashboard"))

    row = None
    try:
        row = (
            supabase.table("uscis_submissions")
            .select("*")
            .eq("submission_id", sid)
            .single()
            .execute()
        ).data
    except Exception as e:
        print("ℹ️ uscis_submissions lookup failed:", e)

    if not row:
        try:
            row = (
                supabase.table("submissions")
                .select("*")
                .eq("submission_id", sid)
                .single()
                .execute()
            ).data
        except Exception as e:
            print("ℹ️ legacy submissions lookup failed:", e)

    if not row:
        return "❌ Submission not found.", 404

    return render_template("grader/nomas_review.html", submission=row)


@lti.route("/nomas/approve-and-send", methods=["POST"])
def nomas_approve_and_send():
    submission_id = request.form.get("submission_id")
    if not submission_id:
        from flask import flash

        flash("Missing submission id.", "error")
        return redirect(url_for("lti.nomas_dashboard"))

    try:
        supabase.table("uscis_submissions").update(
            {"reviewed": True, "pending": False}
        ).eq("submission_id", submission_id).execute()
        # TODO: If you need LMS passback for NoMas, call a NoMas-specific post-back here.
        from flask import flash

        flash("Submission approved.", "success")
    except Exception as e:
        current_app.logger.exception("Approve error: %s", e)
        from flask import flash

        flash("Failed to approve.", "error")

    return redirect(url_for("lti.nomas_dashboard"))


@lti.route("/nomas/delete-submission", methods=["POST"])
def nomas_delete_submission():
    submission_id = request.form.get("submission_id") or (
        request.is_json and (request.get_json() or {}).get("submission_id")
    )
    if not submission_id:
        return jsonify({"success": False, "error": "Missing submission_id"}), 400

    try:
        supabase.table("uscis_submissions").delete().eq(
            "submission_id", str(submission_id)
        ).execute()
        if request.is_json:
            return jsonify({"success": True}), 200
        from flask import flash

        flash("Submission deleted.", "success")
    except Exception as e:
        current_app.logger.exception("Delete error: %s", e)
        if request.is_json:
            return jsonify({"success": False, "error": "Delete failed"}), 500
        from flask import flash

        flash("Failed to delete submission.", "error")
    return redirect(url_for("lti.nomas_dashboard"))


@lti.route("/nomas/save-notes", methods=["POST"], endpoint="nomas_save_notes")
def nomas_save_notes():
    """
    Save instructor notes for NoMas/USCIS submissions.
    Writes to `uscis_submissions` and stays on the same page using next_url.
    """
    submission_id = (request.form.get("submission_id") or "").strip()
    new_notes = request.form.get("instructor_notes", "")
    next_url = request.form.get("next_url")  # e.g. "/nomas/review?submission_id=..."

    if not submission_id:
        return "❌ Submission ID missing", 400

    # --- RLS: set client UID from the row's student_id if available ---
    uid = None
    try:
        rec = (
            supabase.table("uscis_submissions")
            .select("student_id")
            .eq("submission_id", submission_id)
            .single()
            .execute()
        )
        if rec.data:
            uid = str(rec.data.get("student_id") or "")
    except Exception as e:
        print("⚠️ uscis_submissions lookup failed in save-notes:", str(e))

    if not uid:
        uid = str(session.get("student_id") or session.get("user_id") or "")

    if uid:
        try:
            supabase.rpc("set_client_uid", {"uid": uid}).execute()
            print("🔐 set_client_uid:", uid)
        except Exception as e:
            print("⚠️ set_client_uid RPC failed (continuing):", str(e))

    # --- Update notes ---
    try:
        upd = (
            supabase.table("uscis_submissions")
            .update({"instructor_notes": new_notes})
            .eq("submission_id", submission_id)
            .execute()
        )
        if not (upd.data and len(upd.data) > 0):
            print("⚠️ Notes update affected 0 rows; submission_id:", submission_id)
            return "❌ Could not save notes (permission or row not found).", 403

        print("✅ Instructor notes saved for:", submission_id)
    except Exception as e:
        print("❌ Error updating uscis_submissions notes:", str(e))
        return f"❌ Error saving notes: {e}", 500

    # --- Go back to where the form asked us to ---
    if next_url:
        return redirect(next_url)

    # Safe fallback: back to the review page for this submission
    return redirect(url_for("lti.nomas_review", submission_id=submission_id))


@lti.route("/post-grade", methods=["POST"])
def post_grade():
    score = int(request.form.get("score", 0))
    feedback = request.form.get("feedback", "")
    launch_data = session.get("launch_data", {})
    ags_claim = launch_data.get(
        "https://purl.imsglobal.org/spec/lti-ags/claim/endpoint"
    )

    if ags_claim and "lineitem" in ags_claim:
        try:
            lineitem_url = ags_claim["lineitem"].split("?")[0] + "/scores"
            private_key_path = os.path.join("app", "keys", "private_key.pem")
            with open(private_key_path) as f:
                private_key = f.read()

            oauth = OAuth1Session(
                client_key=os.getenv("CLIENT_ID"),
                signature_method="RSA-SHA1",
                rsa_key=private_key,
                signature_type="auth_header",
            )

            assignment_title = request.form.get("assignment_title", "").strip()
            assignment_config = load_assignment_config(assignment_title)

            score_payload = {
                "userId": launch_data.get("sub"),
                "scoreGiven": score,
                "scoreMaximum": assignment_config.get("total_points"),
                "activityProgress": "Completed",
                "gradingProgress": "FullyGraded",
                "timestamp": datetime.utcnow().isoformat() + "Z",
            }

            ags_response = oauth.post(
                lineitem_url,
                json=score_payload,
                headers={"Content-Type": "application/vnd.ims.lis.v1.score+json"},
            )
            ags_response.raise_for_status()
            print("✅ Manual grade posted by instructor.")
        except Exception as e:
            print("❌ Instructor grade post failed:", str(e))

    return render_template(
        "feedback.html",
        score=None,
        feedback=None,
        rubric_total_points=None,
        user_roles=session.get("launch_data", {}).get(
            "https://purl.imsglobal.org/spec/lti/claim/roles", []
        ),
        pending_message="✅ Your assignment was submitted successfully! Your grade and personalized feedback will be available after the review window.",
    )


@lti.route("/test-grader", methods=["GET", "POST"])
def test_grader():
    try:
        response = supabase.table("assignments").select("*").execute()
        rubric_index = response.data or []
    except Exception as e:
        print("❌ Supabase fetch error in test-grader:", e)
        flash("❌ Error loading assignments.", "danger")
        rubric_index = []

    selected_config = None
    gpt_prompt = ""
    gpt_feedback = ""
    gpt_score = None

    if request.method == "POST":
        assignment_title = request.form.get("assignment_title")
        submission_text = request.form.get("submission_text", "").strip()

        selected_config = next(
            (
                cfg
                for cfg in rubric_index
                if cfg["assignment_title"] == assignment_title
            ),
            None,
        )

        if not selected_config:
            return "❌ No config found for that assignment.", 400

        # ✅ Check if grading should be delayed
        # ✅ Check if grading should be delayed
        delay_setting = selected_config.get("delay_posting", "immediate")

        delay_map = {"immediate": 0, "12h": 12, "24h": 24, "36h": 36, "48h": 48}

        delay_hours = delay_map.get(delay_setting, 0)

        if delay_hours > 0:
            from datetime import datetime, timedelta

            from app.storage import (
                store_pending_feedback,
            )  # Ensure this import is valid

            release_time = datetime.utcnow() + timedelta(hours=delay_hours)

            store_pending_feedback(
                assignment_title=assignment_title,
                student_id="test_user",  # Replace with real student ID later
                feedback="(To be generated at release time)",
                score=None,
                release_time=release_time.isoformat(),
            )

            gpt_feedback = f"⏳ Feedback for '{assignment_title}' will be generated after {delay_hours} hour(s)."
            gpt_score = None

            return render_template(
                "test_grader.html",
                rubric_index=rubric_index,
                selected_config=selected_config,
                gpt_prompt="",
                gpt_feedback=gpt_feedback,
                gpt_score=gpt_score,
            )

        # Load rubric
        rubric_text = ""
        rubric_path = os.path.join("rubrics", selected_config.get("rubric_file", ""))
        try:
            if rubric_path.endswith(".json"):
                with open(rubric_path, encoding="utf-8") as f:
                    rubric_json = json.load(f)

                if "sections" in rubric_json:
                    rubric_text = "\n".join(
                        [f"- {section['title']}" for section in rubric_json["sections"]]
                    )
                    rubric_total_points = rubric_json.get("total_points")
                elif "criteria" in rubric_json:
                    rubric_text = "\n".join(
                        [f"- {c['description']}" for c in rubric_json["criteria"]]
                    )
                    # rubric_total_points = get_total_points_from_rubric(rubric_json)
                else:
                    return (
                        "❌ Unrecognized rubric format. Please upload a valid .json rubric with 'sections' or 'criteria'.",
                        400,
                    )

            elif rubric_path.endswith(".docx"):
                doc = Document(rubric_path)
                rubric_text = "\n".join([para.text for para in doc.paragraphs])
            elif rubric_path.endswith(".pdf"):
                rubric_text = extract_pdf_text(rubric_path)
        except:
            rubric_text = "(Unable to load rubric.)"

        # Build GPT prompt
        prompt = f"""
You are a helpful AI grader.

Assignment Title: {assignment_title}
Grading Difficulty: {selected_config.get("grading_difficulty")}
Student Level: {selected_config.get("student_level")}
Feedback Tone: {selected_config.get("feedback_tone")}
Total Points: {selected_config.get("total_points")}

Rubric:
{rubric_text}
"""

        if selected_config.get("ai_notes"):
            prompt += f"""

Instructor Notes:
{selected_config["ai_notes"]}
"""

        prompt += f"""

Student Submission:
---
The following is the student's full submission. Please preserve paragraph formatting, line breaks, and indentation when analyzing or quoting their writing.

\"\"\"{submission_text}\"\"\"
---

Return your response in this format:

Score: <number from 0 to {selected_config.get("total_points")}>
Feedback: <detailed, helpful feedback in paragraph form>

(Note: A future version of this tool may request table-based feedback. If no such instruction is present, return standard narrative feedback only.)
"""

        gpt_prompt = prompt.strip()

        try:
            model_to_use = selected_config.get("gpt_model", "gpt-4")
            openai.api_key = os.getenv("OPENAI_API_KEY")
            response = openai.ChatCompletion.create(
                model=model_to_use,
                messages=[{"role": "user", "content": gpt_prompt}],
                temperature=0.5,
                max_tokens=500,
            )
            output = response["choices"][0]["message"]["content"]

            score_match = re.search(r"Score:\s*(\d{1,3})", output)
            gpt_score = int(score_match.group(1)) if score_match else None

            feedback_match = re.search(r"Feedback:\s*(.+)", output, re.DOTALL)
            gpt_feedback = (
                feedback_match.group(1).strip() if feedback_match else output.strip()
            )

            log_gpt_interaction(assignment_title, gpt_prompt, gpt_feedback, gpt_score)

        except Exception as e:
            gpt_feedback = f"❌ GPT error: {str(e)}"

    return render_template(
        "test_grader.html",
        rubric_index=rubric_index,
        selected_config=selected_config,
        gpt_prompt=gpt_prompt,
        gpt_feedback=gpt_feedback,
        gpt_score=gpt_score if "gpt_score" in locals() else None,
    )


@lti.route("/save-assignment", methods=["POST"])
def save_assignment():
    """
    Create/Update a Rubiqs Grader assignment.
    - Ensures rows show up in /grader-base by including tool="grader"
    - Only sends UUID-typed columns when valid (avoids PostgREST type errors)
    - Adds course_id + created_at for better filtering/sorting
    """
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    print("🚨 HIT /save-assignment")

    # 🔐 RLS: tell PostgREST who we are
    apply_rls_uid()

    # -------- Form fields --------
    assignment_title = (request.form.get("assignment_title") or "").strip()
    if not assignment_title:
        return "❌ Assignment title is required", 400
    
    # --- NEW: independent display title + stable slug (H5P-style) ---
    # display_title is what authors see in Rubiqs; slug is the stable ID used from LMS
    display_title = (request.form.get("display_title") or assignment_title).strip()
    user_slug = (request.form.get("slug") or "").strip()
    slug = slugify(user_slug or display_title)

    lms_link_title = (request.form.get("lms_link_title") or display_title).strip()


    grading_difficulty = request.form.get("grading_difficulty")
    grade_level = request.form.get("grade_level")
    total_points_raw = (request.form.get("total_points") or "").strip()
    total_points = int(total_points_raw) if total_points_raw.isdigit() else 0

    gpt_model = request.form.get("gpt_model", "gpt-4")
    requires_review = (
        request.form.get("requires_review", "false").strip().lower() == "true"
    )
    gospel_enabled = (
        request.form.get("gospel_enabled", "false").strip().lower() == "true"
    )
    custom_ai = request.form.get("custom_ai", "")
    allow_inline = (
        request.form.get("allow_inline_submission", "no").strip().lower() == "yes"
    )
    complete_incomplete = (
        request.form.get("complete_incomplete", "false").strip().lower() == "true"
    )

    assignment_type = request.form.get("assignment_type", "essay")
    scenario_intro = request.form.get("scenario_intro", "")
    gpt_persona = request.form.get("gpt_persona", "")
    student_goal = request.form.get("student_goal", "")
    rubric_criteria = request.form.get("rubric_criteria", "")
    rubric_criteria_list = (
        [c.strip() for c in rubric_criteria.split(",") if c.strip()]
        if rubric_criteria
        else []
    )

    rubric_file = request.files.get("rubric_upload")
    answer_key_file = request.files.get("answer_key_upload")
    additional_file = request.files.get("additional_files")

    # -------- File uploads -> storage URLs --------
    from werkzeug.utils import secure_filename

    # Use the assignment title to namespace files in the bucket
    secure_title = secure_filename(assignment_title).replace(" ", "_")

    def _save_and_upload(fileobj, bucket: str) -> str:
        """
        Upload the given FileStorage to Supabase and return a PUBLIC URL string.
        """
        fname = secure_filename(fileobj.filename).replace(" ", "_")

        # storage key inside the bucket: <assignment-title>/<filename>
        key = f"{secure_title}/{fname}" if secure_title else fname

        # read file bytes
        data = fileobj.read()
        # reset pointer in case file is used again later
        try:
            fileobj.seek(0)
        except Exception:
            pass

        # content type if Werkzeug provided one
        content_type = getattr(fileobj, "mimetype", None)

        # upload the bytes using the canonical helper
        upload_to_supabase(bucket, key, data, content_type=content_type)

        # return a public URL string so downstream code keeps working
        return supabase.storage.from_(bucket).get_public_url(key)

    rubric_url = ""
    if rubric_file and rubric_file.filename:
        rubric_url = _save_and_upload(rubric_file, "rubrics") or ""
    answer_key_url = ""
    if answer_key_file and answer_key_file.filename:
        answer_key_url = _save_and_upload(answer_key_file, "rubrics") or ""
    additional_url = ""
    if additional_file and additional_file.filename:
        additional_url = _save_and_upload(additional_file, "attachments") or ""

    # normalize accidental double-folder
    if rubric_url:
        rubric_url = rubric_url.rstrip("?").replace("rubrics/rubrics/", "rubrics/")
    if answer_key_url:
        answer_key_url = answer_key_url.rstrip("?").replace(
            "rubrics/rubrics/", "rubrics/"
        )
    if additional_url:
        additional_url = additional_url.rstrip("?").replace(
            "attachments/attachments/", "attachments/"
        )

    # -------- IDs / scope --------
    assignment_id = f"a_{uuid.uuid4().hex[:8]}"
    course_id = session.get("course_id")  # TEXT in your schema
    created_at_iso = datetime.utcnow().isoformat()

    # Only include UUID-typed columns when valid
    created_by_uuid = _uuid_or_none(session.get("user_id"))
    institution_uuid = _uuid_or_none(session.get("institution_id"))

    # -------- Build payload (includes tool="grader") --------
    payload = {
        "assignment_id": assignment_id,
        "assignment_title": assignment_title,
        "display_title": display_title,
        "slug": slug,
        "lms_link_title": lms_link_title,
        # required for Grader dashboards
        "tool": "grader",
        "course_id": course_id,
        "created_at": created_at_iso,
        # rubric/files
        "rubric_file": rubric_url,
        "answer_key_file": answer_key_url,
        "additional_file": additional_url,
        # scoring/flags
        "total_points": total_points,
        "instructor_approval": requires_review,
        "requires_persona": False,
        "faith_integration": gospel_enabled,
        "grading_difficulty": grading_difficulty,
        "gpt_model": gpt_model,
        "student_level": grade_level,
        "feedback_tone": "supportive",
        "ai_notes": custom_ai,
        "allow_inline_submission": allow_inline,
        "complete_incomplete": complete_incomplete,
        # chat-ish extras (harmless for grader)
        "assignment_type": assignment_type,
        "scenario_intro": scenario_intro,
        "gpt_persona": gpt_persona,
        "student_goal": student_goal,
        "rubric_criteria": rubric_criteria_list,
    }

    # RLS-scoped UUID columns ONLY if valid
    if created_by_uuid:
        payload["created_by"] = created_by_uuid
    if institution_uuid:
        payload["institution_id"] = institution_uuid

    # -------- Insert --------
    try:
        ins = supabase.table("assignments").insert(payload).execute()
        data = getattr(ins, "data", None)
        err = getattr(ins, "error", None)
        print("📝 insert data:", data)
        print("🧯 insert error:", err)
        if not data:
            raise RuntimeError(err or "Insert returned no rows")
    except Exception as e:
        # Print PostgREST body when available
        print("❌ Insert exception:", repr(e))
        resp = getattr(e, "response", None)
        if resp is not None:
            try:
                print("❌ PostgREST text:", getattr(resp, "text", None))
                print("❌ PostgREST json:", resp.json())
            except Exception:
                pass

        msg = str(e)
        if "row-level security" in msg.lower():
            flash(
                "❌ You don’t have permission to save this assignment (RLS).", "danger"
            )
        elif "duplicate key value" in msg.lower():
            flash(
                "❌ Assignment already exists with the same name. Please use a different title.",
                "danger",
            )
        else:
            flash(f"❌ Error saving assignment: {msg}", "danger")
        return redirect(url_for("lti.grader_base"))

    # ✅ Success → back to dashboard HTML
    return redirect(url_for("lti.grader_base", success=assignment_title))


@lti.route("/instructor-review", methods=["GET", "POST"], endpoint="instructor_review")
def instructor_review():
    submission_id = request.values.get("submission_id")  # args or form
    current_review = None
    next_id = None

    print("🔍 is_superuser:", session.get("is_superuser"))
    print("🔍 course_id:", session.get("course_id"))
    print("🔍 institution_id:", session.get("institution_id"))

    # ---------- POST: update score/feedback ----------
    if request.method == "POST":
        print("📩 POST form data:", dict(request.form))
        submission_id = (request.form.get("submission_id") or "").strip()
        if not submission_id:
            return "❌ submission_id required", 400

        # set RLS uid first
        uid = None
        try:
            rec = (
                supabase.table("submissions")
                .select("student_id,user_id")
                .eq("submission_id", submission_id)
                .single()
                .execute()
            )
            if rec.data:
                uid = rec.data.get("student_id") or rec.data.get("user_id")
        except Exception as e:
            print("⚠️ lookup for RLS uid failed:", str(e))

        if uid:
            try:
                supabase.rpc("set_client_uid", {"uid": str(uid)}).execute()
            except Exception as e:
                print("⚠️ set_client_uid RPC failed (continuing):", str(e))

        # do not touch submission_time here
        updated_score = request.form.get("score")
        updated_feedback = request.form.get("feedback", "")

        try:
            supabase.table("submissions").update(
                {
                    "score": int(updated_score) if updated_score is not None else None,
                    "feedback": updated_feedback,
                    "reviewed": True,
                    "pending": False,
                }
            ).eq("submission_id", submission_id).execute()
        except Exception as e:
            print("❌ update failed:", str(e))
            return f"❌ Failed to save review: {e}", 500

        # go back to the same submission page
        return redirect(url_for("lti.instructor_review", submission_id=submission_id))

    # ---------- GET: load one submission or a review queue ----------
    if submission_id:
        resp = (
            supabase.table("submissions")
            .select("*")
            .eq("submission_id", submission_id)
            .execute()
        )
        reviews = resp.data or []
    else:
        if session.get("is_superuser"):
            print("👑 Superuser: all unreviewed")
            resp = (
                supabase.table("submissions")
                .select("*")
                .eq("pending", True)
                .eq("reviewed", False)
                .order("submission_time", desc=True)
                .execute()
            )
        else:
            print("👤 Instructor: filtered unreviewed by institution/course")
            resp = (
                supabase.table("submissions")
                .select("*")
                .eq("institution_id", session.get("institution_id"))
                .eq("course_id", session.get("course_id"))
                .eq("pending", True)
                .eq("reviewed", False)
                .order("submission_time", desc=True)
                .execute()
            )
        reviews = resp.data or []

    # sanity
    reviews = [
        r for r in reviews if r.get("submission_id") and r.get("assignment_title")
    ]

    if submission_id:
        current_review = reviews[0] if reviews else None
    if not current_review and reviews:
        current_review = reviews[0]
        if len(reviews) > 1:
            next_id = reviews[1]["submission_id"]

    # hotfix: hide any seeded test note if it exists
    bad = "Test feedback (RLS check)"
    if current_review and (current_review.get("instructor_notes") or "").strip() == bad:
        current_review["instructor_notes"] = ""

    # template handles the viewer (Google gview iframe)
    return render_template(
        "instructor_review.html",
        current_review=current_review,
        reviews=reviews,
        next_id=next_id,
    )


@lti.route("/instructor-review/save-notes", methods=["POST"])
def instructor_save_notes():
    submission_id = request.form.get("submission_id")
    new_notes = request.form.get("instructor_notes", "")

    if not submission_id:
        return "❌ Submission ID missing", 400

    record = (
        supabase.table("submissions")
        .select("student_id")
        .eq("submission_id", submission_id)
        .single()
        .execute()
    )
    if record.data:
        uid = str(record.data.get("student_id") or session.get("user_id"))
    if uid:
        supabase.rpc("set_client_uid", {"uid": uid}).execute()
        print("🔐 Using set_client_uid with:", uid)

        print("🔐 Using set_client_uid with:", uid)

    response = (
        supabase.table("submissions")
        .update({"instructor_notes": new_notes})
        .eq("submission_id", submission_id)
        .execute()
    )

    print("✅ Instructor notes saved for:", submission_id)
    return redirect("/instructor-review?submission_id=" + submission_id)


@lti.route("/instructor-review-button", methods=["GET", "POST"])
def instructor_review_button():
    print("🔍 is_superuser:", session.get("is_superuser"))
    print("🔍 institution_id:", session.get("institution_id"))
    print("🔍 course_id:", session.get("course_id"))

    if session.get("is_superuser"):
        print("👑 Superuser: loading all unreviewed submissions")
        response = (
            supabase.table("submissions")
            .select("*")
            .eq("pending", True)
            .eq("reviewed", False)
            .execute()
        )
    else:
        print(
            "👤 Instructor: filtering unreviewed submissions by institution and course"
        )
        response = (
            supabase.table("submissions")
            .select("*")
            .eq("institution_id", session.get("institution_id"))
            .eq("course_id", session.get("course_id"))
            .eq("pending", True)
            .eq("reviewed", False)
            .execute()
        )

    reviews = response.data or []

    # ✅ Filter out any malformed entries
    reviews = [
        r for r in reviews if r.get("submission_id") and r.get("assignment_title")
    ]

    print(f"🧪 Number of pending reviews found: {len(reviews)}")  # ✅ Add this

    submission_id = request.args.get("submission_id")

    # ❌ Prevent crash if invalid submission_id is passed
    if submission_id in (None, "", "undefined"):
        return render_template(
            "instructor_review.html",
            current_review=None,
            reviews=[],
            html_output="",
            next_id=None,
        )

    if request.method == "POST":
        submission_id = request.form.get("submission_id")
        updated_score = int(request.form.get("score"))
        updated_feedback = request.form.get("feedback")

        supabase.table("submissions").update(
            {
                "score": updated_score,
                "feedback": updated_feedback,
                "timestamp": datetime.utcnow().isoformat(),
            }
        ).eq("submission_id", submission_id).execute()

        return redirect(url_for("lti.instructor_review_button"))

    current_review = None
    if submission_id:
        current_review = next(
            (r for r in reviews if r["submission_id"] == submission_id), None
        )
    elif reviews:
        current_review = reviews[0]

        return render_template(
            "instructor_review.html", current_review=current_review, reviews=reviews
        )


def post_grade_to_lms(session, score, feedback):
    print("🧪 lineitem_url:", session.get("lineitem_url"))
    print("🧪 feedback:", feedback)

    try:
        launch_data = session.get("launch_data", {})
        ags_claim = launch_data.get(
            "https://purl.imsglobal.org/spec/lti-ags/claim/endpoint"
        )

        if not ags_claim or "lineitem" not in ags_claim:
            print("⚠️ AGS info missing — cannot post grade.")
            return

        lineitem_url = ags_claim["lineitem"].split("?")[0] + "/scores"
        private_key_path = os.path.join("app", "keys", "private_key.pem")
        with open(private_key_path) as f:
            private_key = f.read()

        oauth = OAuth1Session(
            client_key=os.getenv("CLIENT_ID"),
            signature_method="RSA-SHA1",
            rsa_key=private_key,
            signature_type="auth_header",
        )

        assignment_title = (
            session.get("launch_data", {})
            .get("https://purl.imsglobal.org/spec/lti/claim/resource_link", {})
            .get("title", "")
            .strip()
        )
        print("📝 Assignment Title:", assignment_title)
        assignment_config = load_assignment_config(assignment_title)

        if not assignment_config or not assignment_config.get("total_points"):
            print(
                f"❌ Missing assignment config or total_points for: {assignment_title}"
            )
            return

        rubric_total_points = assignment_config.get("total_points")
        if rubric_total_points is None:
            print(f"❌ total_points is None for assignment: {assignment_title}")
            return

        score_payload = {
            "userId": launch_data.get("sub"),
            "scoreGiven": score,
            "scoreMaximum": rubric_total_points,
            "comment": str(feedback or "Graded with Rubiqs"),
            "activityProgress": "Completed",
            "gradingProgress": "FullyGraded",
            "timestamp": datetime.utcnow().isoformat() + "Z",
        }

        print("📦 Final Score Payload:")
        for k, v in score_payload.items():
            print(f"  {k}: {repr(v)} ({type(v)})")

        response = oauth.post(
            lineitem_url,
            json=score_payload,
            headers={"Content-Type": "application/vnd.ims.lis.v1.score+json"},
        )

        if response.status_code >= 200 and response.status_code < 300:
            print("✅ Grade posted to LMS.")
        else:
            print("⚠️ AGS post failed:", response.text)

    except Exception as e:
        print("❌ Error in post_grade_to_lms():", str(e))


@lti.route("/edit-assignment/<assignment_id>", methods=["GET", "POST"])
def edit_assignment(assignment_id):
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    if request.method == "POST":
        print("🚀 Save Assignment POST route hit")
        try:
            print("📥 POST data:", request.form)

            total_points = request.form.get("total_points", type=int)
            gpt_model = request.form.get("gpt_model", "gpt-4")
            ai_notes = request.form.get("ai_notes", "")
            student_level = request.form.get("student_level")
            grading_difficulty = request.form.get("grading_difficulty")
            delay_posting = request.form.get("delay_posting", "immediate")
            allow_inline = request.form.get("allow_inline_submission", "no") == "yes"

            faith_raw = request.form.get("faith_integration", "false")
            faith_integration = faith_raw.lower() == "true"

            print(
                "🧠 Final values to save:",
                {
                    "total_points": total_points,
                    "ai_notes": ai_notes,
                    "student_level": student_level,
                    "grading_difficulty": grading_difficulty,
                    "gpt_model": gpt_model,
                    "faith_integration": faith_integration,
                    "delay_posting": delay_posting,
                    "allow_inline_submission": allow_inline,
                },
            )

            response = (
                supabase.table("assignments")
                .update(
                    {
                        "total_points": total_points,
                        "ai_notes": ai_notes,
                        "student_level": student_level,
                        "grading_difficulty": grading_difficulty,
                        "gpt_model": gpt_model,
                        "faith_integration": faith_integration,
                        "delay_posting": delay_posting,
                        "allow_inline_submission": allow_inline,
                    }
                )
                .eq("assignment_id", assignment_id)
                .execute()
            )

            if hasattr(response, "error") and response.error:
                print("❌ Supabase error:", response.error.message)
                return f"❌ Supabase update error: {response.error.message}", 500

            print("✅ Assignment updated successfully")
            return redirect(url_for("lti.view_assignments"))

        except Exception as e:
            print("❌ Exception in edit_assignment:", e)
            return "Internal Server Error", 500

    # ✅ GET request logic — check in assignments first
    response = (
        supabase.table("assignments")
        .select("*")
        .eq("assignment_id", assignment_id)
        .execute()
    )
    data = response.data

    if not data:
        # ✅ Fallback: Try chat_assignments and redirect if found
        response = (
            supabase.table("chat_assignments")
            .select("*")
            .eq("assignment_id", assignment_id)
            .execute()
        )
        data = response.data
        if data:
            print("↪️ Redirecting to edit_chat_assignment route...")
            return redirect(
                url_for("lti.edit_chat_assignment", assignment_id=assignment_id)
            )
        else:
            return f"No assignment found with ID {assignment_id}", 404

    assignment = data[0]
    return render_template(
        "edit_assignment.html",
        assignment=assignment,
        tinymce_api_key=os.getenv("TINYMCE_API_KEY"),
    )


@lti.route("/grader-edit", methods=["GET"])
def grader_edit_alias():
    assignment_id = request.args.get("assignment_id")
    if not assignment_id:
        return "Missing assignment_id", 400
    return redirect(
        url_for("lti.edit_assignment", assignment_id=assignment_id), code=302
    )


@lti.route("/view-assignments")
def view_assignments():
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    # 🔐 RLS visibility
    apply_rls_uid()

    try:
        course_id = session.get("course_id", "demo_course")

        if session.get("is_superuser"):
            print("👑 Superuser: loading all assignments")
            response = supabase.table("assignments").select("*").execute()
        else:
            print(
                "👤 Instructor: filtering assignments by institution and course (allow legacy NULLs)"
            )
            q = supabase.table("assignments").select("*").eq("tool", "grader")

            inst_id = session.get("institution_id")
            if inst_id:
                q = q.or_(f"institution_id.eq.{inst_id},institution_id.is.null")

            course_id = session.get("course_id")
            if course_id:
                q = q.or_(f"course_id.eq.{course_id},course_id.is.null")

            response = q.order("created_at", desc=True).execute()

        assignments = response.data or []

        # 🧪 Add this debug loop AFTER fetching
        for a in assignments:
            print(
                "🧪 ASSIGNMENT DEBUG:",
                a.get("assignment_title"),
                "| ID:",
                a.get("assignment_id"),
            )

    except Exception as e:
        print("❌ Supabase fetch error (view-assignments):", e)
        flash("❌ Error loading assignments.", "danger")
        assignments = []

    return render_template("view_assignments.html", assignments=assignments)


@lti.route("/delete-assignment", methods=["POST"])
def delete_assignment():
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    data = request.get_json()
    assignment_title = data.get("assignment_title", "").strip()
    print(f"🗑 Deleting assignment: {assignment_title}")

    if not assignment_title:
        return jsonify({"success": False, "error": "Missing assignment title"}), 400

    try:
        response = (
            supabase.table("assignments")
            .delete()
            .eq("assignment_title", assignment_title)
            .execute()
        )

        if hasattr(response, "error") and response.error:
            return jsonify({"success": False, "error": response.error.message}), 500

        if not response.data:
            return jsonify({"success": False, "error": "Assignment not found"}), 404

        return jsonify({"success": True})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@lti.route("/nomas-training-dashboard")
def nomas_training_dashboard():
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    inst_id = session.get("institution_id")
    course_id = session.get("course_id", "demo_course")

    # --- Assignments (USCIS/NoMas) ---
    try:
        if session.get("is_superuser"):
            aresp = (
                supabase.table("uscis_assignments")
                .select(
                    "assignment_id, assignment_title, form_type, institution_id, course_id, created_at"
                )
                .order("created_at", desc=True)
                .execute()
            )
            assignments = aresp.data or []
        else:
            # strict first
            aresp = (
                supabase.table("uscis_assignments")
                .select(
                    "assignment_id, assignment_title, form_type, institution_id, course_id, created_at"
                )
                .eq("institution_id", inst_id)
                .eq("course_id", course_id)
                .order("created_at", desc=True)
                .execute()
            )
            assignments = aresp.data or []
            # include legacy rows with NULLs
            if not assignments:
                aresp = (
                    supabase.table("uscis_assignments")
                    .select(
                        "assignment_id, assignment_title, form_type, institution_id, course_id, created_at"
                    )
                    .or_(f"institution_id.is.null,institution_id.eq.{inst_id}")
                    .or_(f"course_id.is.null,course_id.eq.{course_id}")
                    .order("created_at", desc=True)
                    .execute()
                )
                assignments = aresp.data or []
    except Exception as e:
        print("❌ load uscis_assignments:", e)
        assignments = []

    # Build a quick lookup of title -> form_type for fallback mapping
    assignment_ft_by_title = {
        (a.get("assignment_title") or ""): (a.get("form_type") or None)
        for a in assignments
    }

    # --- Submissions (NoMas only) — tolerate missing columns like institution_id/course_id ---
    try:
        sresp = (
            supabase.table("uscis_submissions")
            .select("*")
            .order("submission_time", desc=True)
            .execute()
        )
        submissions = sresp.data or []
    except Exception as e:
        current_app.logger.exception("❌ Failed to load NoMas submissions: %s", e)
        submissions = []

    # --- Fallback: if uscis_submissions is empty, surface legacy 'submissions' rows that correspond to NoMas assignments ---
    try:
        if not submissions:
            titles = [
                a.get("assignment_title")
                for a in assignments
                if a.get("assignment_title")
            ]
            if titles:
                if session.get("is_superuser"):
                    legacy = (
                        supabase.table("submissions")
                        .select("*")
                        .in_("assignment_title", titles)
                        .order("submission_time", desc=True)
                        .limit(200)
                        .execute()
                    ).data or []
                else:
                    legacy = (
                        supabase.table("submissions")
                        .select("*")
                        .in_("assignment_title", titles)
                        .or_(f"institution_id.is.null,institution_id.eq.{inst_id}")
                        .or_(f"course_id.is.null,course_id.eq.{course_id}")
                        .order("submission_time", desc=True)
                        .limit(200)
                        .execute()
                    ).data or []
                # Map legacy rows to the template shape
                mapped = []
                for r in legacy:
                    mapped.append(
                        {
                            **r,
                            "user_id": r.get("student_id"),
                            "submitted_at": r.get("submission_time"),
                            "form_type": assignment_ft_by_title.get(
                                r.get("assignment_title") or "", None
                            ),
                        }
                    )
                submissions = mapped
    except Exception as e:
        print("⚠️ fallback from submissions failed:", e)

    return render_template(
        "grader/nomas_training_dashboard.html",
        assignments=assignments,
        submissions=submissions,
    )


@lti.route("/_debug/nomas-latest")
def _debug_nomas_latest():
    rows = (
        supabase.table("uscis_submissions")
        .select("*")
        .order("submitted_at", desc=True)
        .limit(5)
        .execute()
        .data
        or []
    )
    return jsonify(rows)


@lti.route("/release-pending", methods=["GET"])
def release_pending_feedback():
    print("🚀 /release-pending triggered")

    try:
        from datetime import datetime

        now = datetime.utcnow().isoformat()

        response = (
            supabase.table("submissions")
            .select("*")
            .eq("pending", True)
            .lt("release_time", now)
            .execute()
        )

        if hasattr(response, "error") and response.error:
            print("❌ Supabase query error:", response.error.message)
            return f"Supabase query failed: {response.error.message}", 500

        pending = response.data or []
        print(f"📬 Found {len(pending)} entries eligible for release")

        released = 0

        for entry in pending:
            assignment_id = entry.get("assignment_id")
            student_id = entry.get("student_id")
            score = entry.get("score")
            feedback = entry.get("feedback")
            submission_id = entry.get("submission_id")

            if not all([assignment_id, student_id, score, feedback, submission_id]):
                print(f"⚠️ Skipping incomplete submission: {submission_id}")
                continue

            if student_id:
                supabase.rpc("set_client_uid", {"uid": str(student_id)}).execute()
                print("🔐 Using set_client_uid with:", student_id)

            update_response = (
                supabase.table("submissions")
                .update({"pending": False, "reviewed": True, "released_at": now})
                .eq("submission_id", submission_id)
                .execute()
            )

            if hasattr(update_response, "error") and update_response.error:
                print(
                    f"❌ Error updating submission {submission_id}: {update_response.error.message}"
                )
                continue

            released += 1

        return f"✅ Released {released} submissions", 200

    except Exception as e:
        print("❌ Fatal error in release process:", str(e))
        return f"❌ Internal error: {str(e)}", 500


@lti.route("/run-delay-checker")
def run_delay_checker():
    from datetime import datetime

    # Fetch submissions that are still waiting
    response = (
        supabase.table("submissions").select("*").eq("ready_to_post", False).execute()
    )

    if not response.data:
        print("✅ No submissions pending release.")
        return "✅ No pending submissions to check.", 200

    now = datetime.utcnow()
    updates_made = 0

    for submission in response.data:
        try:
            # Get the release_time from the database (this is already stored when you save submissions)
            release_time = datetime.fromisoformat(
                submission["release_time"].replace("Z", "")
            )  # Remove 'Z' if present

            # Check if the current time is greater than or equal to the release time
            if now >= release_time:
                # Update ready_to_post to True if delay has expired
                supabase.table("submissions").update({"ready_to_post": True}).eq(
                    "submission_id", submission["submission_id"]
                ).execute()
                updates_made += 1

        except Exception as e:
            print(
                f"❌ Error checking submission {submission.get('submission_id')}: {str(e)}"
            )

    print(f"✅ Delay check complete. Updated {updates_made} submissions.")

    return f"✅ Delay check complete. Updated {updates_made} submissions.", 200


@lti.route(
    "/grader/download-activity-log",
    methods=["GET"],
    endpoint="grader_download_activity_log",
)
def grader_download_activity_log():
    from flask import redirect, url_for

    return redirect(url_for("lti.download_activity_log"))


# 🧪 Triggering redeploy


@lti.route("/delete-submission", methods=["POST"])
def delete_submission():
    try:
        submission_id = request.form.get("submission_id") or request.get_json().get(
            "submission_id"
        )
        print("🧪 DELETE REQUEST RECEIVED:", submission_id)

        if not submission_id:
            print("❌ No submission_id received")
            return (
                jsonify({"success": False, "error": "No submission_id provided"}),
                400,
            )

        parsed_id = str(submission_id)

        # Try to fetch the submission
        record = (
            supabase.table("submissions")
            .select("student_id")
            .eq("submission_id", parsed_id)
            .single()
            .execute()
        )
        if not record.data:
            print("❌ Submission not found in Supabase:", parsed_id)
            return jsonify({"success": False, "error": "Submission not found"}), 404

        uid = str(record.data.get("student_id") or session.get("user_id"))
        if uid:
            supabase.rpc("set_client_uid", {"uid": uid}).execute()
            print("🔐 Using set_client_uid with:", uid)

        # Delete the record
        response = (
            supabase.table("submissions")
            .delete()
            .eq("submission_id", parsed_id)
            .execute()
        )
        print("🧪 DELETE RESPONSE:", response)

        # Double-check if it's gone
        confirm = (
            supabase.table("submissions")
            .select("*")
            .eq("submission_id", parsed_id)
            .execute()
        )
        if confirm.data:
            print("❌ Record still exists after delete.")
            return jsonify({"success": False, "error": "Delete failed"}), 500

        print("✅ Submission deleted:", parsed_id)
        return jsonify({"success": True}), 200

    except Exception as e:
        print("❌ DELETE ERROR:", str(e))
        return jsonify({"success": False, "error": str(e)}), 500


@lti.route("/instructor-review/accept", methods=["POST"])
def accept_submission():
    if request.is_json:
        submission_id = request.get_json().get("submission_id")
    else:
        submission_id = request.form.get("submission_id")

    print("🧪 ACCEPT REQUEST RECEIVED:", submission_id)

    try:
        if not submission_id:
            return jsonify({"success": False, "error": "Missing submission_id"}), 400

        parsed_id = str(submission_id)
        record = (
            supabase.table("submissions")
            .select("student_id")
            .eq("submission_id", parsed_id)
            .single()
            .execute()
        )

        if not record.data:
            print("❌ Submission not found in Supabase:", parsed_id)
            return jsonify({"success": False, "error": "Submission not found"}), 404

        uid = str(record.data.get("student_id") or session.get("user_id"))
        if uid:
            supabase.rpc("set_client_uid", {"uid": uid}).execute()
            print("🔐 Using set_client_uid with:", uid)

        response = (
            supabase.table("submissions")
            .update({"pending": False, "reviewed": True})
            .eq("submission_id", parsed_id)
            .execute()
        )

        print("🧪 ACCEPT RESPONSE:", response)

        if hasattr(response, "data") and not response.data:
            print("⚠️ No rows updated. Possibly already reviewed.")
            return (
                jsonify(
                    {
                        "success": False,
                        "error": "No matching submission found to update.",
                    }
                ),
                404,
            )

        # ✅ Return JSON if called via fetch(), else redirect
        if request.is_json:
            return jsonify({"success": True}), 200
        else:
            return redirect(url_for("lti.instructor_review"))

    except Exception as e:
        print("❌ ACCEPT ERROR:", str(e))
        return jsonify({"success": False, "error": "Internal error"}), 500


@lti.route("/student-demo-direct")
def student_demo():
    title = request.args.get("title", "Sample USCIS Form Demo").strip()
    assignment_config = load_assignment_config(title)

    if not assignment_config:
        return f"❌ Assignment '{title}' not found in Supabase."

    tinymce_api_key = os.getenv("TINYMCE_API_KEY")

    return render_template(
        "student_demo.html",
        assignment_config=assignment_config,
        tinymce_api_key=tinymce_api_key,
    )


@lti.route("/grade-docx-demo", methods=["POST"])
def grade_docx_demo():
    assignment_title = request.args.get("title", "").strip()
    assignment_title = normalize_title(assignment_title)

    assignment_config = load_assignment_config(assignment_title)
    if not assignment_config:
        return f"❌ Assignment '{assignment_title}' not found in Supabase.", 400

    # Inject demo session
    session["student_id"] = "demo_student"
    session["platform"] = "demo"
    session["launch_data"] = {
        "https://purl.imsglobal.org/spec/lti/claim/resource_link": {
            "title": assignment_title
        },
        "https://purl.imsglobal.org/spec/lti/claim/roles": ["Student"],
    }
    session["course_id"] = "demo_course"

    # Forward to grading logic
    return grade_docx()


@lti.route("/grader-submissions")
def grader_submissions():
    # 🔐 Ensure RLS visibility for this request
    apply_rls_uid()

    try:
        inst_id = session.get("institution_id")
        course_id = session.get("course_id")
        is_super = bool(session.get("is_superuser"))

        print(
            "📥 /grader-submissions | super:",
            is_super,
            "| inst:",
            inst_id,
            "| course:",
            course_id,
        )

        q = (
            supabase.table("submissions")
            .select("*")
            .eq("tool", "grader")
            .order("submission_time", desc=True)
            .limit(300)
        )

        # Include rows for this inst/course OR legacy NULLs
        if not is_super:
            if inst_id:
                q = q.or_(f"institution_id.eq.{inst_id},institution_id.is.null")
            if course_id:
                q = q.or_(f"course_id.eq.{course_id},course_id.is.null")

        resp = q.execute()
        if getattr(resp, "error", None):
            print("❌ Supabase error /grader-submissions:", resp.error)
            return jsonify({"error": "DB error: " + str(resp.error)}), 500

        rows = getattr(resp, "data", None) or []
        print(f"📦 returning {len(rows)} row(s)")

        # Normalize for the frontend
        out = []
        for r in rows:
            out.append(
                {
                    "submission_id": r.get("submission_id") or r.get("id"),
                    "user_id": r.get("user_id") or r.get("student_id"),
                    "assignment_title": r.get("assignment_title")
                    or r.get("assignment_id")
                    or "",
                    "created_at": r.get("created_at")
                    or r.get("submitted_at")
                    or r.get("submission_time"),
                    "reviewed": bool(
                        r.get("reviewed") or r.get("instructor_reviewed") or False
                    ),
                    "score": r.get("score") or r.get("instructor_score"),
                }
            )
        return jsonify(out), 200

    except Exception as e:
        import traceback

        print("❌ Error in /grader-submissions:", repr(e))
        traceback.print_exc()
        return jsonify({"error": "Server error while loading submissions"}), 500


@lti.route("/save-submission", methods=["POST"])
def save_submission():
    try:
        payload = request.get_json(silent=True) or {}
        user_id = payload.get("user_id")
        assignment_title = payload.get("assignment_title") or payload.get(
            "assignment_id"
        )

        if not user_id or not assignment_title:
            return (
                jsonify(
                    {"success": False, "error": "Missing user_id or assignment_title"}
                ),
                400,
            )

        row = {
            "tool": "grader",  # required for the GET query
            "user_id": user_id,
            "assignment_title": assignment_title,
            "content_url": payload.get("content_url"),
            "inline_text": payload.get("inline_text"),
            "score": payload.get("score"),
            "reviewed": False,
            "institution_id": session.get("institution_id"),
            "course_id": session.get("course_id"),
            "submission_time": datetime.utcnow().isoformat() + "Z",
        }

        res = supabase.table("submissions").insert(row).execute()
        if getattr(res, "error", None):
            print("❌ Supabase insert error:", res.error)
            return jsonify({"success": False, "error": str(res.error)}), 500

        saved = (res.data or [row])[0]
        return jsonify({"success": True, "submission": saved}), 200

    except Exception as e:
        import traceback

        print("❌ Error /save-submission:", repr(e))
        traceback.print_exc()
        return jsonify({"success": False, "error": "Server error"}), 500


@lti.route("/grader-assignments", methods=["GET"])
def grader_assignments():
    if "launch_data" not in session and not session.get("logged_in"):
        return redirect(url_for("lti.unauthorized"))

    # RLS uid is already set by before_app_request
    inst = session.get("institution_id")
    course = session.get("course_id")

    try:
        q = (
            supabase.table("assignments")
            .select(
                "assignment_id,assignment_title,total_points,created_at,institution_id,created_by,tool,course_id,gpt_model,instructor_approval"
            )
            .eq("tool", "grader")
            .order("created_at", desc=True)
            .limit(300)
        )

        # Only scope for non-superusers — include legacy NULLs so older rows still appear
        if not (session.get("is_superuser") or session.get("role") == "superuser"):
            if inst:
                q = q.or_(f"institution_id.eq.{inst},institution_id.is.null")
            if course:
                q = q.or_(f"course_id.eq.{course},course_id.is.null")

        res = q.execute()
        rows = res.data or []
        print(f"📦 grader-assignments returning {len(rows)} row(s)")
        return jsonify(rows), 200

    except Exception as e:
        print("❌ grader-assignments error:", e)
        return jsonify({"error": "DB error"}), 500


# Redirect to the canonical route in chat.py and avoid endpoint collisions.
@lti.route(
    "/grader/chat/instructor",
    methods=["GET"],
    endpoint="grader_chat_instructor_dashboard",
)
def grader_chat_instructor_dashboard():
    from flask import redirect, url_for

    return redirect(url_for("lti.instructor_chat_dashboard"))


import os
import re
import traceback

# from openai import OpenAI
# client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def _try_extract_json(text: str):
    m = re.search(r"\{.*\}", text, re.S)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except Exception:
        return None


import os
import re

import openai


@lti.route("/download-mapped-fields")
def download_mapped_fields():
    debug_path = os.path.join("data", "last_mapped_fields.json")
    if not os.path.exists(debug_path):
        return "❌ No file found", 404
    return send_file(debug_path, as_attachment=True)
